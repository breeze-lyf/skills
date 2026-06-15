#!/usr/bin/env python3
"""Lightweight PRD consistency harness for meeting-to-prd-workflow."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path


REQUIRED_HEADINGS = [
    "1. 引言",
    "2. 系统概述",
    "3. 功能需求",
    "4. 数据与集成要求",
    "5. 待确认事项",
]

REQUIRED_TERMS = [
    "客户原始表单截图",
    "字段标记",
    "* 为必填",
    "加粗字段",
    "审批/流转",
    "带出",
    "关联",
    "自动计算",
    "待确认",
]

DEFAULT_LOGIC_TERMS = []


def load_docx(path: Path):
    try:
        from docx import Document
    except Exception as exc:
        raise SystemExit(f"python-docx is required: {exc}")
    return Document(path)


def paragraph_text(doc) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def table_text(doc) -> str:
    chunks = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text.strip())
    return "\n".join(chunks)


def check(path: Path, min_images: int, logic_terms: list[str] | None = None, forbidden_terms: list[str] | None = None):
    doc = load_docx(path)
    paras = paragraph_text(doc)
    all_text = "\n".join(paras) + "\n" + table_text(doc)
    h3_pages = [p for p in paras if re.match(r"^3\.\d+\.\d+\s+", p)]
    screenshot_labels = [p for p in paras if p == "客户原始表单截图："]
    approval_in_tables = []
    for idx, table in enumerate(doc.tables, start=1):
        txt = "\n".join(cell.text for row in table.rows for cell in row.cells)
        if any(term in txt for term in ["签署意见", "审批意见", "总经理批准"]) and "审批/流转" not in txt:
            approval_in_tables.append(idx)

    result = {
        "file": str(path),
        "ok": True,
        "stats": {
            "paragraphs": len(paras),
            "tables": len(doc.tables),
            "images": len(doc.inline_shapes),
            "functional_pages": len(h3_pages),
            "screenshot_labels": len(screenshot_labels),
        },
        "missing_required_headings": [h for h in REQUIRED_HEADINGS if h not in all_text],
        "missing_required_terms": [t for t in REQUIRED_TERMS if t not in all_text],
        "missing_logic_terms": [t for t in (logic_terms or DEFAULT_LOGIC_TERMS) if t not in all_text],
        "forbidden_terms_found": [t for t in (forbidden_terms or []) if t in all_text],
        "approval_like_tables": approval_in_tables,
        "warnings": [],
    }

    if len(doc.inline_shapes) < min_images:
        result["warnings"].append(f"image count {len(doc.inline_shapes)} is below min-images {min_images}")
    if len(h3_pages) and len(screenshot_labels) < max(1, int(len(h3_pages) * 0.6)):
        result["warnings"].append("few screenshot labels relative to functional pages")
    if approval_in_tables:
        result["warnings"].append("approval/signature text may still be inside data tables")

    result["ok"] = not (
        result["missing_required_headings"]
        or result["missing_required_terms"]
        or result["missing_logic_terms"]
        or result["forbidden_terms_found"]
        or result["warnings"]
    )
    return result


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--min-images", type=int, default=1)
    parser.add_argument(
        "--logic-terms",
        default="",
        help="Comma-separated project-specific business terms that must appear, e.g. 门店,巡检任务,地图,AI分析",
    )
    parser.add_argument(
        "--forbidden-terms",
        default="",
        help="Comma-separated out-of-scope terms that must not appear, e.g. BOM,采购订单,开票",
    )
    parser.add_argument("--json", action="store_true")
    args = parser.parse_args()
    logic_terms = [t.strip() for t in args.logic_terms.split(",") if t.strip()]
    forbidden_terms = [t.strip() for t in args.forbidden_terms.split(",") if t.strip()]
    result = check(args.docx, args.min_images, logic_terms, forbidden_terms)
    if args.json:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print(json.dumps(result, ensure_ascii=False, indent=2))
    sys.exit(0 if result["ok"] else 1)


if __name__ == "__main__":
    main()
